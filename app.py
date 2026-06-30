import streamlit as st
import re
import io
from datetime import date
from docx import Document
from docx.oxml.ns import qn

# =============================================================================
# PAGE CONFIG
# =============================================================================

st.set_page_config(
    page_title="Generador de Contratos",
    page_icon="📄",
    layout="centered"
)

st.markdown("""
<style>
    /* Clean, professional form styling */
    .main { background-color: #f7f8fc; }
    section[data-testid="stSidebar"] { display: none; }

    .block-container {
        max-width: 680px;
        padding-top: 2.5rem;
        padding-bottom: 3rem;
    }

    .header-box {
        background: linear-gradient(135deg, #1e3a5f 0%, #2d5f8a 100%);
        border-radius: 12px;
        padding: 28px 32px;
        margin-bottom: 28px;
        color: white;
    }
    .header-box h1 {
        margin: 0 0 6px 0;
        font-size: 22px;
        font-weight: 700;
        letter-spacing: -0.3px;
    }
    .header-box p {
        margin: 0;
        font-size: 14px;
        opacity: 0.8;
    }

    .section-card {
        background: white;
        border-radius: 10px;
        padding: 22px 26px;
        margin-bottom: 16px;
        border: 1px solid #e8eaf0;
    }
    .section-title {
        font-size: 13px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.6px;
        color: #1e3a5f;
        margin-bottom: 14px;
        padding-bottom: 10px;
        border-bottom: 2px solid #e8eaf0;
    }

    div[data-testid="stRadio"] > label { font-weight: 600; }

    .stTextInput > label, .stDateInput > label {
        font-size: 13px !important;
        font-weight: 600 !important;
        color: #374151 !important;
    }
    .stTextInput input, .stDateInput input {
        border-radius: 7px !important;
        border-color: #d1d5db !important;
        font-size: 14px !important;
    }
    .stTextInput input:focus, .stDateInput input:focus {
        border-color: #2d5f8a !important;
        box-shadow: 0 0 0 3px rgba(45,95,138,0.12) !important;
    }

    div[data-testid="stRadio"] div[role="radiogroup"] {
        display: flex;
        gap: 12px;
    }
    div[data-testid="stRadio"] div[role="radiogroup"] label {
        background: #f0f4f8;
        border: 2px solid #d1d5db;
        border-radius: 8px;
        padding: 10px 20px;
        cursor: pointer;
        font-size: 14px;
        font-weight: 500;
        transition: all 0.15s;
    }
    div[data-testid="stRadio"] div[role="radiogroup"] label:has(input:checked) {
        background: #e8f0f8;
        border-color: #2d5f8a;
        color: #1e3a5f;
    }

    .stButton > button {
        background: linear-gradient(135deg, #1e3a5f 0%, #2d5f8a 100%);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 12px 32px;
        font-size: 15px;
        font-weight: 600;
        width: 100%;
        cursor: pointer;
        transition: opacity 0.15s;
        margin-top: 8px;
    }
    .stButton > button:hover { opacity: 0.9; }

    .stDownloadButton > button {
        background: #16a34a !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 12px 32px !important;
        font-size: 15px !important;
        font-weight: 600 !important;
        width: 100% !important;
    }

    .info-note {
        background: #eff6ff;
        border-left: 4px solid #2d5f8a;
        border-radius: 0 6px 6px 0;
        padding: 10px 14px;
        font-size: 13px;
        color: #1e3a5f;
        margin-top: 12px;
    }

    .error-box {
        background: #fef2f2;
        border-left: 4px solid #dc2626;
        border-radius: 0 6px 6px 0;
        padding: 10px 14px;
        font-size: 13px;
        color: #dc2626;
        margin-top: 8px;
    }

    .success-box {
        background: #f0fdf4;
        border: 1px solid #86efac;
        border-radius: 8px;
        padding: 16px 20px;
        text-align: center;
        font-size: 14px;
        color: #15803d;
        font-weight: 600;
        margin-bottom: 12px;
    }
</style>
""", unsafe_allow_html=True)


# =============================================================================
# CONTRACT LOGIC (identical to Colab version)
# =============================================================================

ORDINALS_ES = [
    'PRIMERA', 'SEGUNDA', 'TERCERA', 'CUARTA', 'QUINTA',
    'SEXTA', 'SÉPTIMA', 'OCTAVA', 'NOVENA', 'DÉCIMA',
    'UNDÉCIMA', 'DUODÉCIMA', 'DECIMOTERCERA', 'DECIMOCUARTA', 'DECIMOQUINTA',
    'DECIMOSEXTA', 'DECIMOSÉPTIMA', 'DECIMOCTAVA', 'DECIMONOVENA', 'VIGÉSIMA',
]

def merge_runs_in_para(para):
    if len(para.runs) <= 1:
        return
    full_text = ''.join(r.text for r in para.runs)
    para.runs[0].text = full_text
    for run in para.runs[1:]:
        run.text = ''

def get_paragraph_text(para):
    merge_runs_in_para(para)
    return para.runs[0].text.strip() if para.runs else ''

def paragraph_is_clause_heading(text):
    upper = text.upper().strip()
    for pre in ['CLÁUSULAS ', 'CLAUSULAS ']:
        if upper.startswith(pre):
            upper = upper[len(pre):]
            break
    for word in ORDINALS_ES:
        if upper.startswith(word):
            return word
    return None

def find_clause_block(doc, target_ordinal):
    paras = doc.paragraphs
    start = None
    for i, para in enumerate(paras):
        text = get_paragraph_text(para)
        ordinal = paragraph_is_clause_heading(text)
        if ordinal and ordinal == target_ordinal.upper():
            start = i
        elif ordinal and start is not None:
            return start, i - 1
    if start is not None:
        return start, len(paras) - 1
    return None, None

def remove_paragraphs_by_index(doc, indices_to_remove):
    body = doc.element.body
    paras = doc.paragraphs
    elements_to_remove = [paras[i]._element for i in indices_to_remove if i < len(paras)]
    for el in elements_to_remove:
        body.remove(el)

def is_antiguedad_heading(text):
    upper = text.upper().strip()
    return upper.startswith('SEGUNDA') and ('ANTIGÜ' in upper or 'ANTIGU' in upper)

def is_numbered_clause(text):
    upper = text.upper().strip()
    check = upper
    for pre in ['CLÁUSULAS ', 'CLAUSULAS ']:
        if upper.startswith(pre):
            check = upper[len(pre):]
            break
    for word in ORDINALS_ES:
        if check.startswith(word):
            idx = upper.find(word)
            return word, idx
    return None, None

def renumber_clauses(doc):
    for para in doc.paragraphs:
        merge_runs_in_para(para)
    heading_paras = []
    for para in doc.paragraphs:
        full_text = get_paragraph_text(para)
        ordinal, idx = is_numbered_clause(full_text)
        if ordinal is not None:
            heading_paras.append((para, ordinal, idx))
    subpoint_map = {}
    for i, (para, detected, idx) in enumerate(heading_paras):
        correct_num = i + 1
        if detected in ORDINALS_ES:
            labelled_as = ORDINALS_ES.index(detected) + 1
            if labelled_as != correct_num:
                subpoint_map[labelled_as] = correct_num
    for i, (para, detected, idx) in enumerate(heading_paras):
        correct = ORDINALS_ES[i] if i < len(ORDINALS_ES) else f'CLÁUSULA {i+1}'
        if not para.runs:
            continue
        text = para.runs[0].text
        para.runs[0].text = text[:idx] + correct + text[idx + len(detected):]
    for para in doc.paragraphs:
        if not para.runs:
            continue
        text = para.runs[0].text
        for old_n, new_n in sorted(subpoint_map.items(), reverse=True):
            text = re.sub(
                rf'(^|[\s])({old_n})(\.(\d))',
                lambda m, nn=new_n: f'{m.group(1)}{nn}{m.group(3)}',
                text
            )
        para.runs[0].text = text

def format_date_es(d):
    months = ['', 'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
              'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
    return f"{d.day} de {months[d.month]} de {d.year}"

def format_date_numeric(d):
    return f"{d.day:02d}/{d.month:02d}/{d.year}"

def replace_placeholders(doc, replacements):
    def replace_in_para(para):
        merge_runs_in_para(para)
        if not para.runs:
            return
        text = para.runs[0].text
        for placeholder, value in replacements.items():
            if placeholder in text:
                text = text.replace(placeholder, value)
        para.runs[0].text = text
    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

def remove_note_paragraphs(doc, phrases):
    body = doc.element.body
    to_remove = []
    for para in doc.paragraphs:
        text = ''.join(r.text for r in para.runs).strip().lower()
        if any(phrase.lower() in text for phrase in phrases):
            to_remove.append(para._element)
    for el in to_remove:
        body.remove(el)
    return len(to_remove)

def generate_contract(template_bytes, ctype, name, id_num, nationality, address,
                      contract_date_val, seniority_service_date_val,
                      seniority_indemnity_date_val, external_start_date_val):
    doc = Document(io.BytesIO(template_bytes))

    if ctype == 'internal':
        start, end = find_clause_block(doc, 'CUARTA')
        if start is not None:
            remove_paragraphs_by_index(doc, list(range(start, end + 1)))
        for para in doc.paragraphs:
            merge_runs_in_para(para)
            full = get_paragraph_text(para)
            if is_antiguedad_heading(full):
                para.runs[0].text = re.sub(r'\s*\(si aplica\)', '', para.runs[0].text, flags=re.IGNORECASE).strip()
                break
    else:
        paras = doc.paragraphs
        start = None
        for i, para in enumerate(paras):
            merge_runs_in_para(para)
            text = get_paragraph_text(para)
            if is_antiguedad_heading(text):
                start = i
            elif start is not None and is_numbered_clause(text)[0] is not None:
                remove_paragraphs_by_index(doc, list(range(start, i)))
                break
        else:
            if start is not None:
                remove_paragraphs_by_index(doc, list(range(start, len(doc.paragraphs))))

    renumber_clauses(doc)

    if ctype == 'internal':
        in_puesto = False
        subpoint_counter = 0
        for para in doc.paragraphs:
            merge_runs_in_para(para)
            text = get_paragraph_text(para)
            ordinal, idx = is_numbered_clause(text)
            if ordinal:
                upper = text.upper()
                in_puesto = 'PUESTO' in upper or 'FUNCIONES' in upper
                subpoint_counter = 0
                continue
            if not in_puesto:
                continue
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                continue
            numPr = pPr.find(qn('w:numPr'))
            if numPr is None:
                continue
            pPr.remove(numPr)
            subpoint_counter += 1
            current_text = para.runs[0].text if para.runs else ''
            if para.runs:
                para.runs[0].text = f'3.{subpoint_counter}  {current_text}'
            else:
                para.add_run(f'3.{subpoint_counter}  ')

    c_date = format_date_es(contract_date_val)
    c_date_num = format_date_numeric(contract_date_val)

    for para in doc.paragraphs:
        merge_runs_in_para(para)

    start_date_str = (format_date_numeric(external_start_date_val) if ctype == 'external'
                      else format_date_numeric(seniority_service_date_val) if ctype == 'internal' else '')

    for para in doc.paragraphs:
        if not para.runs:
            continue
        text = para.runs[0].text
        if 'XXX' not in text and 'XXXXXX' not in text:
            continue
        if 'En Madrid' in text or ('a del XXX' in text):
            text = text.replace('XXX', c_date, 1)
        elif 'de nacionalidad' in text or 'con DNI' in text or 'domicilio en' in text:
            text = text.replace('XXX', name, 1)
            text = text.replace('XXX', nationality, 1)
            text = text.replace('XXX', id_num, 1)
            text = text.replace('XXX', address, 1)
        elif 'en fecha XXX' in text or ('Empresa ha propuesto' in text and 'XXX' in text):
            text = text.replace('XXX', c_date_num, 1)
        elif 'prestación de servicios' in text or 'antigüedad es reconocida' in text or ('desde el XXX' in text):
            if ctype == 'internal':
                svc_date = format_date_numeric(seniority_service_date_val)
                ind_date = format_date_numeric(seniority_indemnity_date_val)
                text = text.replace('XXX', svc_date, 1)
                text = text.replace('XXX', ind_date, 1)
        elif 'XXXXXX' in text or 'a partir del día' in text:
            text = text.replace('XXXXXX', start_date_str, 1)
            text = text.replace('XXX', start_date_str, 1)
        para.runs[0].text = text

    remove_note_paragraphs(doc, [
        'cambiar numeración si aplica',
        '(cambiar numeración si aplica)',
        '(cambiar numeración)',
        'cambiar numeración',
    ])

    for para in doc.paragraphs:
        merge_runs_in_para(para)
        if not para.runs:
            continue
        if '(si aplica)' in para.runs[0].text.lower():
            para.runs[0].text = re.sub(r'\s*\(si aplica\)', '', para.runs[0].text, flags=re.IGNORECASE).strip()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


# =============================================================================
# UI
# =============================================================================

st.markdown("""
<div class="header-box">
    <h1>📄 Generador de Contratos</h1>
    <p>Sube la plantilla, rellena los datos y descarga el contrato listo.</p>
</div>
""", unsafe_allow_html=True)

# --- Step 1: Upload template ---
st.markdown('<div class="section-card"><div class="section-title">1 · Plantilla del contrato</div>', unsafe_allow_html=True)
uploaded_file = st.file_uploader("Sube tu plantilla .docx", type=["docx"], label_visibility="collapsed")
if uploaded_file:
    st.success(f"✅ Plantilla cargada: **{uploaded_file.name}**")
else:
    st.markdown('<div class="info-note">📂 Sube el archivo .docx con los marcadores XXX antes de continuar.</div>', unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

# --- Step 2: Contract type ---
st.markdown('<div class="section-card"><div class="section-title">2 · Tipo de contrato</div>', unsafe_allow_html=True)
contract_type = st.radio(
    "Selecciona el tipo:",
    options=["Transferencia interna", "Contratación externa"],
    horizontal=True,
    label_visibility="collapsed"
)
is_internal = contract_type == "Transferencia interna"
st.markdown('</div>', unsafe_allow_html=True)

# --- Step 3: Employee details ---
st.markdown('<div class="section-card"><div class="section-title">3 · Datos del empleado</div>', unsafe_allow_html=True)
col1, col2 = st.columns(2)
with col1:
    employee_name = st.text_input("Nombre completo", placeholder="Ana García López")
with col2:
    employee_id = st.text_input("DNI", placeholder="12345678A")
col3, col4 = st.columns(2)
with col3:
    employee_nationality = st.text_input("Nacionalidad", placeholder="Española")
with col4:
    employee_address = st.text_input("Domicilio", placeholder="Calle Mayor 10, 28001 Madrid")
st.markdown('</div>', unsafe_allow_html=True)

# --- Step 4: Dates ---
st.markdown('<div class="section-card"><div class="section-title">4 · Fechas</div>', unsafe_allow_html=True)
contract_date = st.date_input("Fecha del contrato", value=date.today())

if is_internal:
    st.markdown("**Antigüedad** *(transferencia interna)*")
    col5, col6 = st.columns(2)
    with col5:
        seniority_service_date = st.date_input("Fecha inicio de servicios", value=date.today(), key="svc")
    with col6:
        seniority_indemnity_date = st.date_input("Antigüedad (indemnizatoria)", value=date.today(), key="ind")
    external_start_date = None
else:
    st.markdown("**Fecha de incorporación** *(contratación externa)*")
    external_start_date = st.date_input("Fecha de inicio", value=date.today(), key="ext")
    seniority_service_date = None
    seniority_indemnity_date = None

st.markdown('</div>', unsafe_allow_html=True)

# --- Generate button ---
if st.button("⚙️ Generar contrato"):
    errors = []
    if not uploaded_file:
        errors.append("Sube la plantilla .docx antes de continuar.")
    if not employee_name.strip():
        errors.append("El nombre del empleado es obligatorio.")
    if not employee_id.strip():
        errors.append("El DNI es obligatorio.")
    if not employee_nationality.strip():
        errors.append("La nacionalidad es obligatoria.")
    if not employee_address.strip():
        errors.append("El domicilio es obligatorio.")
    if is_internal and not seniority_service_date:
        errors.append("La fecha de inicio de servicios es obligatoria para transferencias internas.")
    if not is_internal and not external_start_date:
        errors.append("La fecha de inicio es obligatoria para contrataciones externas.")

    if errors:
        for e in errors:
            st.markdown(f'<div class="error-box">❌ {e}</div>', unsafe_allow_html=True)
    else:
        with st.spinner("Generando contrato..."):
            try:
                template_bytes = uploaded_file.read()
                ctype = 'internal' if is_internal else 'external'
                result_bytes = generate_contract(
                    template_bytes=template_bytes,
                    ctype=ctype,
                    name=employee_name.strip(),
                    id_num=employee_id.strip(),
                    nationality=employee_nationality.strip(),
                    address=employee_address.strip(),
                    contract_date_val=contract_date,
                    seniority_service_date_val=seniority_service_date,
                    seniority_indemnity_date_val=seniority_indemnity_date,
                    external_start_date_val=external_start_date,
                )
                type_label = 'interno' if is_internal else 'externo'
                safe_name = re.sub(r'[^a-zA-Z0-9_]', '_', employee_name.strip())
                filename = f"contrato_{type_label}_{safe_name}.docx"

                st.markdown('<div class="success-box">✅ Contrato generado correctamente. Haz clic para descargar.</div>', unsafe_allow_html=True)
                st.download_button(
                    label=f"⬇️ Descargar {filename}",
                    data=result_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Error al generar el contrato: {e}")
